"""Module D2 - session_plan_builder: render one detailed SESSION PLAN to .docx.

Replicates the RVNP/KTTC "SESSION PLAN" template exactly:

  * PORTRAIT US-Letter (21.59 x 27.94 cm), 1.27 cm margins, Times New Roman.
  * Title "SESSION PLAN" (14 pt, bold, centred).
  * ONE 6-column table (100% width, single black grid borders):
      - paired half-rows (Date/Time, Trainer name/number, Institution/Level,
        Class/Number of Trainees),
      - full-width rows (Unit Code, Unit of Competence, Session title, LLN needs,
        Learning outcome(s), Resources, Safety requirements),
      - grey-shaded (D9D9D9) 14 pt section banners: Session Presentation /
        1. Introduction / 2. Session Delivery / 3. Session Review,
      - the delivery sub-grid (Time | Trainer Activity | Trainee Activity |
        Learning Check/Assessment) with one row per teaching step,
      - Assignment, TOTAL TIME and a tall Session Reflection box.
  * Signature block (PREPARED BY / VERIFIED BY / APPROVED BY) double-spaced,
    each with dotted leaders + DATE + SIGN.

Every list renders LINE-BY-LINE (one paragraph per string) - never iterate a raw
string, which is what causes the 'G / r / o / u / p' char-splitting bug.
"""

from __future__ import annotations

import re
from typing import Callable, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE,
                             WD_TABLE_ALIGNMENT)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from models import SessionPlan

FONT_NAME = "Times New Roman"
FONT_SIZE = 12.0           # body / table font size (pt)
BANNER_SIZE = 14.0         # grey section banners + the title's section labels
TITLE_SIZE = 14.0
BANNER_FILL = "D9D9D9"     # grey shading on the section banners
GUIDE_RED = "C00000"       # red used only for the blank-box reflection hint
SIGN_SIZE = 10.0           # bottom sign-off row font size

# 6-column grid (cm) - lifted verbatim from the sample template.
GRID_WIDTHS = [2.52, 2.67, 4.28, 0.73, 4.95, 4.28]
# Dotted leaders for the ROLE / DATE / SIGN fill-in blanks (one sign-off per row).
SIGN_LEADER = "…" * 12
SIGN_LEADER_SHORT = "…" * 8


# --------------------------------------------------------------------------- #
# Low-level run / cell helpers
# --------------------------------------------------------------------------- #
def _style_run(run, size: float = FONT_SIZE, bold: bool = False,
               italic: bool = False, color: Optional[str] = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def _set_cell_shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _clear_paragraphs(cell) -> None:
    cell.text = ""
    # keep exactly the one default paragraph for re-use


def _write_lines(cell, lines: List[str], *, size: float = FONT_SIZE,
                 bold_predicate: Optional[Callable[[str], bool]] = None,
                 all_bold: bool = False, color: Optional[str] = None,
                 align=None,
                 valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER) -> None:
    """Write `lines` into a cell, one paragraph per line (never char-split)."""
    if isinstance(lines, str):
        lines = [lines]
    lines = [str(l) for l in (lines or []) if str(l).strip() != ""] or [""]
    _clear_paragraphs(cell)
    first = cell.paragraphs[0]
    for idx, line in enumerate(lines):
        par = first if idx == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        if align is not None:
            par.alignment = align
        bold = all_bold or bool(bold_predicate and bold_predicate(line))
        _style_run(par.add_run(line), size=size, bold=bold, color=color)
    if valign is not None:
        cell.vertical_alignment = valign


def _write_label_value(cell, label: str, value: str, *,
                       size: float = FONT_SIZE,
                       align=WD_ALIGN_PARAGRAPH.LEFT,
                       valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER) -> None:
    """One line: '<bold label>: <normal value>' (value optional)."""
    _clear_paragraphs(cell)
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.space_before = Pt(0)
    par.alignment = align
    _style_run(par.add_run(label if label.endswith(":") else f"{label}:"),
               size=size, bold=True)
    if str(value).strip():
        _style_run(par.add_run(f" {value}"), size=size, bold=False)
    if valign is not None:
        cell.vertical_alignment = valign


def _write_label_block(cell, label: str, lines: List[str], *,
                       size: float = FONT_SIZE,
                       bold_predicate: Optional[Callable[[str], bool]] = None,
                       bullet: bool = False,
                       valign=WD_CELL_VERTICAL_ALIGNMENT.TOP) -> None:
    """Bold label on the first line, then each value line as its own paragraph.

    With `bullet=True`, each value line renders as a '•' bullet (any leading
    '-'/'•' the source already carried is normalised away first).
    """
    if isinstance(lines, str):
        lines = [lines]
    lines = [str(l) for l in (lines or []) if str(l).strip() != ""]
    _clear_paragraphs(cell)
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(0)
    first.paragraph_format.space_before = Pt(0)
    _style_run(first.add_run(label if label.endswith(":") else f"{label}:"),
               size=size, bold=True)
    for line in lines:
        par = cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        bold = bool(bold_predicate and bold_predicate(line))
        if bullet:
            line = "•  " + re.sub(r"^\s*[-•]\s*", "", line)
        _style_run(par.add_run(line), size=size, bold=bold)
    if valign is not None:
        cell.vertical_alignment = valign


# --------------------------------------------------------------------------- #
# Bold predicates (which whole lines render bold inside a list cell)
# --------------------------------------------------------------------------- #
_GROUP_WORDS = {"knowledge", "skills", "skill", "attitudes", "attitude",
                "knowledge checks"}


def _role_or_group_bold(line: str) -> bool:
    """Bold the 'Trainer:' / 'Trainee(s):' role headers and assessment groups."""
    s = line.strip()
    if s.endswith(":") and len(s.split()) <= 2:
        return True
    return s.lower() in _GROUP_WORDS


_ROLE_PREFIX = re.compile(r"^\s*(trainer|trainee)(\(s\)|s)?\s*[:\-]\s*", re.I)


def _activity_lines(lines: List[str]) -> List[str]:
    """Strip a leading 'Trainer:'/'Trainee(s):' prefix off each activity line.

    The section already carries ONE bold 'Trainer:' label, so a per-line prefix
    (which the model sometimes repeats on every bullet) is redundant and dropped.
    A bare 'Trainer:' line collapses to nothing and is skipped entirely.
    """
    out: List[str] = []
    for raw in (lines or []):
        s = _ROLE_PREFIX.sub("", str(raw).strip()).strip()
        if s:
            out.append(s)
    return out


# --------------------------------------------------------------------------- #
# Table / grid plumbing
# --------------------------------------------------------------------------- #
def _set_grid_cols(table, widths_cm: List[float]) -> None:
    grid = table._element.find(qn("w:tblGrid"))
    for gc in list(grid.findall(qn("w:gridCol"))):
        grid.remove(gc)
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(round(Cm(w).pt * 20))))   # dxa
        grid.append(gc)


def _set_table_width_pct(table, pct: int = 100) -> None:
    tblPr = table._element.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), str(int(pct * 50)))


def _set_cell_margins(table, top: float, bottom: float, left: float,
                      right: float) -> None:
    tblPr = table._element.tblPr
    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        anchor = None
        for tag in ("w:tblLook", "w:tblCaption", "w:tblDescription"):
            anchor = tblPr.find(qn(tag))
            if anchor is not None:
                break
        if anchor is not None:
            anchor.addprevious(mar)
        else:
            tblPr.append(mar)
    for tag, cm in (("w:top", top), ("w:left", left),
                    ("w:bottom", bottom), ("w:right", right)):
        el = mar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            mar.append(el)
        el.set(qn("w:w"), str(int(round(Cm(cm).pt * 20))))
        el.set(qn("w:type"), "dxa")


def _new_row(table, height_cm: Optional[float] = None):
    """Add a fresh 6-cell row, seed cell widths from the grid, return its cells."""
    row = table.add_row()
    if height_cm is not None:
        row.height = Cm(height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell, w in zip(row.cells, GRID_WIDTHS):
        cell.width = Cm(w)
    return row.cells


def _merge(cells, start: int, end: int):
    """Merge grid columns [start, end] inclusive; return the merged cell."""
    merged = cells[start]
    for i in range(start + 1, end + 1):
        merged = merged.merge(cells[i])
    return merged


# --------------------------------------------------------------------------- #
# Row builders
# --------------------------------------------------------------------------- #
def _two_half_rows(table, label_l, value_l, label_r, value_r) -> None:
    cells = _new_row(table)
    left = _merge(cells, 0, 2)
    right = _merge(cells, 3, 5)
    _write_label_value(left, label_l, value_l)
    _write_label_value(right, label_r, value_r)


def _fullwidth_label_value(table, label, value, *,
                           align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cells = _new_row(table)
    merged = _merge(cells, 0, 5)
    _write_label_value(merged, label, value, align=align)


def _fullwidth_block(table, label, lines, *, bold_predicate=None,
                     bullet=False, height_cm=None) -> None:
    cells = _new_row(table, height_cm=height_cm)
    merged = _merge(cells, 0, 5)
    _write_label_block(merged, label, lines, bold_predicate=bold_predicate,
                       bullet=bullet)


def _banner(table, text: str, *, align=None) -> None:
    cells = _new_row(table, height_cm=1.0)
    merged = _merge(cells, 0, 5)
    _set_cell_shade(merged, BANNER_FILL)
    _write_lines(merged, [text], size=BANNER_SIZE, all_bold=True, align=align)


def _delivery_header(table) -> None:
    cells = _new_row(table)
    time_c = _merge(cells, 0, 1)
    trainer_c = _merge(cells, 2, 3)
    trainee_c = cells[4]
    check_c = cells[5]
    _write_lines(time_c, ["Time (in minutes)"], all_bold=True)
    _write_lines(trainer_c, ["Trainer Activity"], all_bold=True)
    _write_lines(trainee_c, ["Trainee Activity"], all_bold=True)
    _write_lines(check_c, ["Learning Check/Assessment"], all_bold=True)


def _delivery_step_row(table, step) -> None:
    cells = _new_row(table)
    label_c = cells[0]
    minutes_c = cells[1]
    trainer_c = _merge(cells, 2, 3)
    trainee_c = cells[4]
    check_c = cells[5]
    _write_lines(label_c, [step.step_label], all_bold=True)
    _write_lines(minutes_c, [f"{step.minutes} minutes"], all_bold=True)
    _write_lines(trainer_c, step.trainer_activity, bold_predicate=_role_or_group_bold,
                 valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    _write_lines(trainee_c, step.trainee_activity, bold_predicate=_role_or_group_bold,
                 valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    _write_lines(check_c, step.learning_check, bold_predicate=_role_or_group_bold,
                 valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)


def _signature_paragraph(doc, role: str) -> None:
    """One sign-off row: 'ROLE: … DATE: … SIGN: …' - all 10 pt, only labels bold."""
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 2.0
    par.paragraph_format.space_after = Pt(0)
    _style_run(par.add_run(f"{role}: "), size=SIGN_SIZE, bold=True)
    _style_run(par.add_run(f"{SIGN_LEADER} "), size=SIGN_SIZE, bold=False)
    _style_run(par.add_run("DATE: "), size=SIGN_SIZE, bold=True)
    _style_run(par.add_run(f"{SIGN_LEADER_SHORT} "), size=SIGN_SIZE, bold=False)
    _style_run(par.add_run("SIGN: "), size=SIGN_SIZE, bold=True)
    _style_run(par.add_run(SIGN_LEADER_SHORT), size=SIGN_SIZE, bold=False)


# --------------------------------------------------------------------------- #
# Page setup + assembly
# --------------------------------------------------------------------------- #
def _setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(21.59)        # US Letter (matches the sample)
    sec.page_height = Cm(27.94)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.27))
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def build_session_plan_document(plan: SessionPlan) -> Document:
    doc = Document()
    _setup_page(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(title.add_run("SESSION PLAN"), size=TITLE_SIZE, bold=True)

    table = doc.add_table(rows=0, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_grid_cols(table, GRID_WIDTHS)
    _set_table_width_pct(table, 100)
    _set_cell_margins(table, top=0.1, bottom=0.1, left=0.15, right=0.15)

    # ----- header identity ------------------------------------------------- #
    _two_half_rows(table, "Date", plan.session_date, "Time", plan.session_time)
    _two_half_rows(table, "Trainer name", plan.trainer_name,
                   "Trainer number", plan.trainer_number)
    _two_half_rows(table, "Institution", plan.institution, "Level", plan.level)
    _two_half_rows(table, "Class", plan.class_code,
                   "Number of Trainees", plan.num_trainees)
    _fullwidth_label_value(table, "Unit Code", plan.unit_code)
    _fullwidth_label_value(table, "Unit of Competence", plan.unit_title)
    _fullwidth_label_value(table, "Session title", plan.session_title)
    _fullwidth_block(
        table,
        "Language, Literacy or Numeracy needs (LLN) or Other Requirements of "
        "learner (group)",
        [plan.lln_requirements] if plan.lln_requirements else [])
    _fullwidth_block(table, "Learning outcome(s)", plan.learning_outcomes)
    _fullwidth_block(table, "Resources (references, and learning aids)",
                     plan.resources)
    _fullwidth_label_value(table, "Safety requirements",
                           plan.safety_requirements)

    # ----- session presentation ------------------------------------------- #
    _banner(table, "Session Presentation", align=WD_ALIGN_PARAGRAPH.CENTER)
    _banner(table, "1.  Introduction (5 minutes)")
    _fullwidth_block(table, "Trainer:", _activity_lines(plan.introduction),
                     bullet=True)

    _banner(table, "2. Session Delivery")
    _delivery_header(table)
    for step in plan.delivery_steps:
        _delivery_step_row(table, step)

    _banner(table, "3.  Session Review: (5 minutes)")
    _fullwidth_block(table, "Trainer:", _activity_lines(plan.review),
                     bullet=True)

    # ----- assignment / total / reflection -------------------------------- #
    cells = _new_row(table)
    merged = _merge(cells, 0, 5)
    _write_label_value(merged, "Assignment", plan.assignment,
                       valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)

    _fullwidth_label_value(table, "TOTAL TIME",
                           f"{plan.total_minutes} minutes",
                           align=WD_ALIGN_PARAGRAPH.CENTER)

    cells = _new_row(table, height_cm=3.5)
    merged = _merge(cells, 0, 5)
    _clear_paragraphs(merged)
    par = merged.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    _style_run(par.add_run("Session Reflection:"), bold=True)
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # ----- signatures ------------------------------------------------------ #
    doc.add_paragraph()
    _signature_paragraph(doc, "PREPARED BY")
    _signature_paragraph(doc, "VERIFIED BY")
    _signature_paragraph(doc, "APPROVED BY")
    return doc


def document_to_bytes(plan: SessionPlan) -> bytes:
    import io
    doc = build_session_plan_document(plan)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_document(plan: SessionPlan, path: str) -> str:
    build_session_plan_document(plan).save(path)
    return path
