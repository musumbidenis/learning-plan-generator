"""Module D - doc_builder: render the Learning Plan to a .docx (RVNP format).

A4 LANDSCAPE, Times New Roman, "Table Grid". ONE combined table that autofits
the page width:
  * header-info rows (paired label/value cells + merged Skill/Benchmark rows),
    laid onto the 9-column grid via cell merges, then
  * the 9-column session grid (REF KTTC/TP/LP/F07).

Every list is rendered LINE-BY-LINE (each string -> its own paragraph). We never
iterate a raw string, which is what causes the 'G / r / o / u / p' char-splitting
bug when a string is mistakenly treated as a list.
"""

from __future__ import annotations

import re
from typing import Callable, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE,
                             WD_TABLE_ALIGNMENT)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from models import PlanInputs, Session, Unit

FONT_NAME = "Times New Roman"
FONT_SIZE = 12.0          # body / table font size (pt) for the whole plan
HEADING_ROW_HEIGHT = 0.8  # cm - the 5 label/value rows of the heading section

SESSION_HEADERS = [
    "Week", "Session No", "Session Title", "Specific Learning Outcomes",
    "Learning Key Points", "Trainee Activities", "Resources & References",
    "Learning Checks / Assessments", "Reflections & Date",
]
# Relative column widths (sum scaled to the usable page width).
SESSION_WIDTHS = [0.55, 0.7, 1.5, 2.1, 2.3, 2.4, 1.8, 1.9, 1.2]


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #
def _style_run(run, size: float = FONT_SIZE, bold: bool = False) -> None:
    """Apply the plan font (Times New Roman, incl. East-Asian mapping) to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    if bold:
        run.bold = True


def _set_cell_text(cell, lines: List[str], bold_predicate: Optional[Callable] = None,
                   size: float = FONT_SIZE, header: bool = False) -> None:
    """Write `lines` into a table cell, one paragraph per line.

    `bold_predicate(line) -> bool` decides which whole lines render bold.
    A bare string is wrapped to a single-element list to defeat char-splitting.
    """
    if isinstance(lines, str):
        lines = [lines]
    lines = [l for l in (lines or []) if str(l).strip() != ""] or [""]

    # reuse the default first paragraph, then add the rest
    cell.text = ""
    first_par = cell.paragraphs[0]
    for idx, line in enumerate(lines):
        par = first_par if idx == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run = par.add_run(str(line))
        bold = header or bool(bold_predicate and bold_predicate(str(line)))
        _style_run(run, size=size, bold=bold)


def _is_caps_heading(line: str) -> bool:
    """A 'mostly uppercase' content sub-heading (e.g. OVERVIEW OF ...)."""
    letters = [c for c in line if c.isalpha()]
    if len(letters) < 3:
        return False
    upper = sum(1 for c in letters if c.isupper())
    return upper / len(letters) >= 0.85 and not line.strip().startswith("-")


def _keypoint_bold(line: str) -> bool:
    return _is_caps_heading(line)


def _activity_bold(line: str) -> bool:
    s = line.strip().lower()
    return s.startswith("follow up activity") or s.startswith("due date")


def _assessment_bold(line: str) -> bool:
    s = line.strip().lower()
    return s.startswith(("knowledge checks", "skills:", "skills :", "attitudes"))


def _outcome_bold(line: str) -> bool:
    return line.strip().lower().startswith("by the end")


def _set_table_cell_margins(table, top: float, bottom: float,
                            left: float, right: float) -> None:
    """Uniform interior padding (cm) for EVERY cell in the table.

    Written as a table-level <w:tblCellMar>, inserted in its correct schema
    slot (before tblLook/tblCaption/tblDescription) so Word never flags the
    document for repair.
    """
    tblPr = table._element.tblPr
    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        anchor = None
        for tag in ("w:tblLook", "w:tblCaption", "w:tblDescription"):
            anchor = tblPr.find(qn(tag))
            if anchor is not None:
                break
        if anchor is not None:
            anchor.addprevious(mar)
        else:
            tblPr.append(mar)
    for tag, cm in (("w:top", top), ("w:left", left),
                    ("w:bottom", bottom), ("w:right", right)):
        el = mar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            mar.append(el)
        el.set(qn("w:w"), str(int(round(Cm(cm).pt * 20))))   # dxa = twentieths of a pt
        el.set(qn("w:type"), "dxa")


def _set_table_width_pct(table, pct: int = 100) -> None:
    """Force the table preferred width to `pct`% of the page (AutoFit to Window)."""
    tblPr = table._element.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), str(int(pct * 50)))   # width in fiftieths of a percent


def _autofit_widths(table, widths_cm: List[float]) -> None:
    """AutoFit the table to the page width while keeping column proportions.

    `table.autofit = True` lets Word recompute column widths; the preferred
    cell widths (set only on full 9-cell grid rows) seed the proportions, and a
    100%-of-page table width stretches the grid to fill the landscape page.
    """
    table.autofit = True
    _set_table_width_pct(table, 100)
    for row in table.rows:
        if len(row.cells) != len(widths_cm):
            continue   # skip merged header rows; the grid columns drive their span
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


# --------------------------------------------------------------------------- #
# Header info rows (laid onto the shared 9-column grid via cell merges)
# --------------------------------------------------------------------------- #
def _set_label_value_cell(cell, label: str, value: str) -> None:
    """Render '<bold label>: <value>' as one left-aligned, vertically-centred line."""
    cell.text = ""                       # collapse a merged cell's empty paragraphs
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.space_before = Pt(0)
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_run(par.add_run(f"{label}:"), bold=True)
    if str(value).strip():
        _style_run(par.add_run(f" {value}"), bold=False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_two_col_heading_row(table, label_l, value_l, label_r, value_r) -> None:
    """A heading row of TWO cells, each holding a bold 'Label: value' pair.

    Left cell spans grid columns 0-4, right cell spans 5-8 (~half and half).
    """
    row = table.add_row()
    # Keep the heading rows a neat, uniform height (grows only if a value wraps).
    row.height = Cm(HEADING_ROW_HEIGHT)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    c = row.cells                        # 9 grid cells
    left = c[0]
    for cell in c[1:5]:
        left = left.merge(cell)
    right = c[5]
    for cell in c[6:9]:
        right = right.merge(cell)
    _set_label_value_cell(left, label_l, value_l)
    _set_label_value_cell(right, label_r, value_r)


def _add_fullwidth_row(table, lines, bold_predicate) -> None:
    """A single row whose 9 grid cells are merged into one full-width cell."""
    c = table.add_row().cells
    merged = c[0]
    for cell in c[1:]:
        merged = merged.merge(cell)
    _set_cell_text(merged, lines, size=FONT_SIZE, bold_predicate=bold_predicate)


def _skill_task_lines(unit: Unit) -> List[str]:
    """'Skill or Job Task' = the OS Unit Description from its first action verb.

    Falls back to the numbered element titles only when no description was parsed.
    """
    if unit.skill_task.strip():
        return [unit.skill_task]
    return [f"{el.number}. {el.title}" for el in unit.elements] or [""]


def _benchmark_lines(unit: Unit) -> List[str]:
    """Elements (1,2,3) each followed by their PCs keeping x.y numbering."""
    out: List[str] = []
    for el in unit.elements:
        out.append(f"{el.number}. {el.title}")
        for pc in el.performance_criteria:
            out.append(f"{pc.number} {pc.text}")
    return out


def add_header_rows(table, unit: Unit, inputs: PlanInputs,
                    display_code: str) -> None:
    """Append the header-info rows to the shared 9-column session table."""
    _add_two_col_heading_row(table, "Unit of Competence", unit.unit_title,
                             "Unit Code", display_code)
    _add_two_col_heading_row(table, "Institution", inputs.institution,
                             "Name of Trainer", inputs.trainer_name)
    _add_two_col_heading_row(table, "Course", inputs.course,
                             "Level", inputs.level or unit.level)
    _add_two_col_heading_row(table, "Date of Preparation", inputs.date_of_preparation,
                             "Date of Revision", "")
    _add_two_col_heading_row(table, "Number of Trainees", str(inputs.num_trainees),
                             "Class", inputs.class_code)

    # Merged 'Skill or Job Task' row
    _add_fullwidth_row(
        table, ["Skill or Job Task:"] + _skill_task_lines(unit),
        bold_predicate=lambda l: l.strip().lower().startswith("skill or job"))

    # Merged 'Benchmark or Criteria' row
    _add_fullwidth_row(
        table, ["Benchmark or Criteria to be used:"] + _benchmark_lines(unit),
        bold_predicate=lambda l: l.strip().lower().startswith("benchmark or criteria")
        or bool(re.match(r"^\d+\.\s+\D", l.strip())))


# --------------------------------------------------------------------------- #
# Session table
# --------------------------------------------------------------------------- #
def add_session_grid(table, sessions: List[Session]) -> None:
    """Append the column-header row and one row per session to the table."""
    head_cells = table.add_row().cells
    for cell, head in zip(head_cells, SESSION_HEADERS):
        _set_cell_text(cell, [head], size=FONT_SIZE, header=True)

    for s in sessions:
        cells = table.add_row().cells
        _set_cell_text(cells[0], [str(s.week)], size=FONT_SIZE)
        _set_cell_text(cells[1], [s.session_no], size=FONT_SIZE)
        _set_cell_text(cells[2], [s.session_title], size=FONT_SIZE)
        _set_cell_text(cells[3], s.learning_outcomes, size=FONT_SIZE,
                       bold_predicate=_outcome_bold)
        _set_cell_text(cells[4], s.key_points, size=FONT_SIZE, bold_predicate=_keypoint_bold)
        _set_cell_text(cells[5], s.trainee_activities, size=FONT_SIZE,
                       bold_predicate=_activity_bold)
        _set_cell_text(cells[6], s.resources, size=FONT_SIZE)
        _set_cell_text(cells[7], s.assessments, size=FONT_SIZE,
                       bold_predicate=_assessment_bold)
        _set_cell_text(cells[8], [""], size=FONT_SIZE)   # Reflections & Date - blank


# --------------------------------------------------------------------------- #
# Document assembly
# --------------------------------------------------------------------------- #
def _setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    # A4 landscape
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.0))
    # default style font
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def build_document(unit: Unit, sessions: List[Session], inputs: PlanInputs,
                   display_code: str = "") -> Document:
    doc = Document()
    _setup_page(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("LEARNING PLAN")
    trun.bold = True
    trun.font.name = FONT_NAME
    trun.font.size = Pt(14)

    # One combined table: header-info rows first, then the 9-column session grid.
    table = doc.add_table(rows=0, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_rows(table, unit, inputs, display_code or unit.os_code)
    add_session_grid(table, sessions)

    # Breathing room inside every cell so the grid reads cleanly.
    _set_table_cell_margins(table, top=0.1, bottom=0.1, left=0.15, right=0.15)

    # AutoFit the single table to the landscape page width.
    total = sum(SESSION_WIDTHS)
    usable = 27.7   # cm (A4 landscape, ~1cm margins)
    widths = [w / total * usable for w in SESSION_WIDTHS]
    _autofit_widths(table, widths)
    return doc


def save_document(unit: Unit, sessions: List[Session], inputs: PlanInputs,
                  path: str, display_code: str = "") -> str:
    doc = build_document(unit, sessions, inputs, display_code)
    doc.save(path)
    return path


def document_to_bytes(unit: Unit, sessions: List[Session], inputs: PlanInputs,
                      display_code: str = "") -> bytes:
    import io
    doc = build_document(unit, sessions, inputs, display_code)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
