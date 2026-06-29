"""Module D - doc_builder tests. ZERO API calls."""

import io

from docx import Document

import ai_client
import doc_builder
import planner
from models import PlanInputs


def _build(prog_curr_unit, prog_os_unit):
    inp = PlanInputs(trainer_name="T", institution="RVNP", level="6",
                     num_trainees="25", class_code="L6ICT",
                     date_of_preparation="05-05-2026",
                     term_weeks=13, cat_weeks=[4, 8, 13])
    sessions = planner.plan_sessions(prog_curr_unit, prog_os_unit, inp)
    sessions = ai_client.merge_ai_into_sessions(sessions, [], prog_os_unit)
    data = doc_builder.document_to_bytes(prog_os_unit, sessions, inp,
                                         display_code=prog_curr_unit.curriculum_code)
    return Document(io.BytesIO(data)), sessions


# Combined table layout: 5 label/value rows + 2 full-width rows = 7 header-info
# rows (indices 0-6), then the session column-header row (index 7), then one row
# per session (indices 8+).
_SESSION_HEADER_ROW = 7
_FIRST_SESSION_ROW = 8


def test_single_combined_table(prog_curr_unit, prog_os_unit):
    doc, sessions = _build(prog_curr_unit, prog_os_unit)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    header = [c.text for c in table.rows[_SESSION_HEADER_ROW].cells]
    assert header == doc_builder.SESSION_HEADERS
    # session rows == planned sessions
    assert len(table.rows) - _FIRST_SESSION_ROW == len(sessions)


def test_header_rows_values(prog_curr_unit, prog_os_unit):
    doc, _ = _build(prog_curr_unit, prog_os_unit)
    rows = doc.tables[0].rows
    # value1 spans grid cols 3-4, value2 spans cols 7-8 (merged cells repeat)
    assert rows[0].cells[3].text == "APPLY COMPUTER PROGRAMMING PRINCIPLES"
    assert rows[0].cells[7].text == "IT/CU/ICTA/CC/02/5/MA"
    # Skill/Job task row lists the element titles
    assert "Apply computer programming skills" in rows[5].cells[0].text
    # Benchmark row keeps PC numbering
    assert "1.1 Programming language types" in rows[6].cells[0].text


def test_no_char_splitting_in_cells(prog_curr_unit, prog_os_unit):
    doc, _ = _build(prog_curr_unit, prog_os_unit)
    activities_cell = doc.tables[0].rows[_FIRST_SESSION_ROW].cells[5]
    # each activity is its own paragraph; none is a single character
    paras = [p.text for p in activities_cell.paragraphs if p.text]
    assert len(paras) >= 3
    assert all(len(p) > 1 for p in paras)


def test_reflections_column_blank(prog_curr_unit, prog_os_unit):
    doc, _ = _build(prog_curr_unit, prog_os_unit)
    for row in doc.tables[0].rows[_FIRST_SESSION_ROW:]:
        assert row.cells[8].text.strip() == ""
