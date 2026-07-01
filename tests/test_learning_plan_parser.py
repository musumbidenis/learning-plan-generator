"""Round-trip: build a Learning Plan .docx, then read it back into data.

ZERO API calls - doc_builder + learning_plan_parser are both pure Python. The
parser must recover the header identity and every session so a user can upload a
Learning Plan and generate Session Plans without re-uploading OS + Curriculum.
"""

import os
import tempfile

import doc_builder
import learning_plan_parser as lpp
import session_plan_builder
from models import (Element, PerformanceCriterion, PlanInputs, Session, Unit)


def _unit() -> Unit:
    return Unit(
        unit_title="Apply Market Research", os_code="HOS/OS/03", level="5",
        skill_task="Apply market research in a hospitality establishment.",
        elements=[
            Element(number="1", title="Identify market research tools",
                    performance_criteria=[
                        PerformanceCriterion("1.1", "Market research tools are identified"),
                        PerformanceCriterion("1.2", "Tools are selected appropriately")]),
            Element(number="2", title="Conduct market research",
                    performance_criteria=[
                        PerformanceCriterion("2.1", "Data is collected")]),
        ])


def _sessions():
    return [
        Session(week=1, session_no="1", is_cat=False,
                session_title="Market Research Tools and Techniques",
                pcs=["1.1 Market research tools are identified"],
                key_points=["IDENTIFICATION OF TOOLS", "- Questionnaires",
                            "- Interviews"],
                learning_outcomes=["By the end of the session, the trainee should "
                                   "be able to:", "a. Define market research tools.",
                                   "b. Identify at least four tools."],
                trainee_activities=["- Group Discussion: explore tools.",
                                    "Follow up Activity:",
                                    "1. Develop two tools. (15 Marks)", "Due date:"],
                resources=["- Kotler, P. (2021). Marketing Management. Pearson.",
                           "- PowerPoint on market research"],
                assessments=["Knowledge Checks:", "1. Oral questioning",
                             "Skills:", "1. Observation of developed tools"]),
        Session(week=4, session_no="1", is_cat=True,
                session_title="CAT 1",
                pcs=[], key_points=["ASSESSMENT COVERAGE"],
                learning_outcomes=["By the end of the session, the trainee should "
                                   "be able to:", "a. Demonstrate competence."],
                trainee_activities=["- Complete the Continous Assessment Test(CAT)."],
                resources=["Assessment Tool(s)", "Assessor Guide"],
                assessments=["-Graded Knowledge."]),
    ]


def _inputs() -> PlanInputs:
    return PlanInputs(
        trainer_name="Jane Doe", institution="The Rift Valley National Polytechnic",
        course="ICT Technician", level="5", num_trainees="25",
        class_code="HM-2A", date_of_preparation="01-07-2026")


def _roundtrip(tmp_path):
    unit, sessions, inputs = _unit(), _sessions(), _inputs()
    path = os.path.join(tmp_path, "lp.docx")
    doc_builder.save_document(unit, sessions, inputs, path, display_code="HOS/OS/03")
    return lpp.parse_learning_plan(path)


def test_roundtrip_recovers_header_identity(tmp_path):
    unit, sessions, inputs = _roundtrip(tmp_path)
    assert unit.unit_title == "Apply Market Research"
    assert unit.os_code == "HOS/OS/03"
    assert unit.level == "5"
    assert inputs.trainer_name == "Jane Doe"
    assert inputs.institution == "The Rift Valley National Polytechnic"
    assert inputs.course == "ICT Technician"
    assert inputs.num_trainees == "25"
    assert inputs.class_code == "HM-2A"
    assert inputs.date_of_preparation == "01-07-2026"


def test_roundtrip_recovers_every_session(tmp_path):
    _, sessions, _ = _roundtrip(tmp_path)
    assert len(sessions) == 2
    s0 = sessions[0]
    assert s0.week == 1 and s0.session_no == "1"
    assert s0.session_title == "Market Research Tools and Techniques"
    assert s0.learning_outcomes[0].startswith("By the end")
    assert any("Kotler" in r for r in s0.resources)
    assert "IDENTIFICATION OF TOOLS" in s0.key_points
    assert any("Develop two tools" in a for a in s0.trainee_activities)


def test_roundtrip_flags_cat_session(tmp_path):
    _, sessions, _ = _roundtrip(tmp_path)
    assert sessions[1].is_cat is True
    assert sessions[0].is_cat is False


def test_roundtrip_recovers_benchmark_elements(tmp_path):
    unit, _, _ = _roundtrip(tmp_path)
    assert [e.number for e in unit.elements] == ["1", "2"]
    assert unit.elements[0].title == "Identify market research tools"
    pcs = {pc.number: pc.text for e in unit.elements for pc in e.performance_criteria}
    assert pcs["1.1"] == "Market research tools are identified"
    assert pcs["2.1"] == "Data is collected"


def test_parsed_plan_feeds_session_plan_generation(tmp_path):
    """The recovered data must be usable by the offline session-plan builder."""
    import ai_client
    unit, sessions, inputs = _roundtrip(tmp_path)
    plan = ai_client.build_session_plan_offline(
        unit, sessions[0], inputs, display_code=unit.os_code, duration_minutes=120)
    assert plan.session_title == "Market Research Tools and Techniques"
    assert plan.delivery_steps
    assert plan.unit_code == "HOS/OS/03"


def test_non_learning_plan_docx_raises(tmp_path):
    from docx import Document
    path = os.path.join(tmp_path, "not_a_plan.docx")
    Document().add_paragraph("Just some notes, no table.")
    Document().save(path)
    try:
        lpp.parse_learning_plan(path)
        assert False, "expected LearningPlanParseError"
    except lpp.LearningPlanParseError:
        pass


# --------------------------------------------------------------------------- #
# Batch zip packaging
# --------------------------------------------------------------------------- #
def test_plans_to_zip_packs_all_documents(tmp_path):
    import io
    import zipfile
    import ai_client
    unit, sessions, inputs = _unit(), _sessions(), _inputs()
    named = []
    for s in sessions:
        plan = ai_client.build_session_plan_offline(unit, s, inputs,
                                                    display_code=unit.os_code)
        named.append((f"Session_Plan_W{s.week}_S{s.session_no}.docx", plan))
    data = session_plan_builder.plans_to_zip(named)
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = zf.namelist()
    assert len(names) == 2
    assert all(n.endswith(".docx") for n in names)


def test_plans_to_zip_disambiguates_duplicate_names(tmp_path):
    import io
    import zipfile
    import ai_client
    unit, sessions, inputs = _unit(), _sessions(), _inputs()
    plan = ai_client.build_session_plan_offline(unit, sessions[0], inputs)
    data = session_plan_builder.plans_to_zip(
        [("dup.docx", plan), ("dup.docx", plan)])
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = zf.namelist()
    assert len(set(names)) == 2                    # no silent overwrite
