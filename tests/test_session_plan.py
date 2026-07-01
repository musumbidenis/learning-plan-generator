"""Session Plan: deterministic builder, AI safety nets, and .docx rendering.

ZERO API calls - the AI merge is tested by feeding canned dicts straight into the
coercion/merge helpers, and the offline builder is fully deterministic.
"""

from docx.oxml.ns import qn
from docx.shared import Pt

import ai_client
import session_plan_builder
from models import PlanInputs, Session, SessionPlan, Unit


def _unit() -> Unit:
    return Unit(unit_title="Apply Market Research", os_code="HOS/OS/03", level="5",
                assessment_methods=["Oral questioning", "Observation",
                                    "Portfolio of evidence"])


def _session() -> Session:
    return Session(
        week=2, session_no="1", is_cat=False,
        session_title="Market Research Tools and Techniques",
        pcs=["1.1 Market research tools are identified"],
        key_points=["IDENTIFICATION OF TOOLS", "- Questionnaires"],
        learning_outcomes=["By the end of the session, the trainee should be able to:",
                           "a. Define market research tools.",
                           "b. Identify at least four tools."],
        trainee_activities=["- Group Discussion: explore tools in groups.",
                            "- Think-Pair-Share: compare tools.",
                            "- Demonstration with Participation: build a questionnaire.",
                            "Follow up Activity:",
                            "1. Develop two tools for an establishment. (15 Marks)",
                            "Due date:"],
        resources=["- Kotler, P. (2021). Marketing Management. Pearson.",
                   "- PowerPoint on market research"],
        assessments=["Knowledge Checks:", "1. Oral questioning", "2. Written assessment",
                     "Skills:", "1. Observation of developed tools",
                     "Attitudes:", "1. Attention to detail"])


def _inputs() -> PlanInputs:
    return PlanInputs(trainer_name="Jane Doe", institution="RVNP", level="5",
                      num_trainees="25", class_code="HM-2A")


# --------------------------------------------------------------------------- #
# Offline (deterministic) builder
# --------------------------------------------------------------------------- #
def test_offline_plan_is_complete_and_grounded():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), display_code="HOS/OS/03",
        trainer_number="TR/9", session_date="23/8/2026", duration_minutes=120)

    assert plan.unit_code == "HOS/OS/03"
    assert plan.session_title == "Market Research Tools and Techniques"
    assert plan.trainer_name == "Jane Doe"
    assert plan.introduction and plan.review
    assert plan.delivery_steps
    # carried over verbatim from the Learning-Plan session
    assert plan.learning_outcomes[0].startswith("By the end")
    assert any("Kotler" in r for r in plan.resources)
    # assignment is taken from the session's follow-up activity
    assert "Develop two tools" in plan.assignment
    assert plan.lln_requirements and plan.safety_requirements


def test_offline_step_minutes_sum_to_total():
    duration = 120
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=duration)
    assert sum(s.minutes for s in plan.delivery_steps) == duration - 10
    assert plan.total_minutes == duration          # 5' intro + steps + 5' review


def test_offline_practice_step_includes_skills_check():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    last = plan.delivery_steps[-1].learning_check
    assert any(line.strip().lower() == "skills" for line in last)


# --------------------------------------------------------------------------- #
# Assessment grouping
# --------------------------------------------------------------------------- #
def test_group_assessments_splits_by_heading():
    groups = ai_client._group_assessments(_session().assessments)
    assert groups["Knowledge"] == ["1. Oral questioning", "2. Written assessment"]
    assert groups["Skills"] == ["1. Observation of developed tools"]
    assert groups["Attitudes"] == ["1. Attention to detail"]


# --------------------------------------------------------------------------- #
# AI merge safety nets
# --------------------------------------------------------------------------- #
def test_merge_empty_ai_falls_back_to_defaults():
    body = ai_client._merge_session_plan_body({}, _unit(), _session(), 120)
    assert body["delivery_steps"]                  # never empty
    assert body["introduction"] and body["review"]
    assert body["assignment"]


def test_merge_keeps_valid_ai_steps():
    ai_rows = {
        "introduction": ["Trainer:", "Takes roll call."],
        "delivery_steps": [
            {"step_label": "Step 1", "minutes": 40,
             "trainer_activity": ["Trainer:", "Leads a Group Discussion."],
             "trainee_activity": ["Trainee(s):", "Discuss tools."],
             "learning_check": ["Knowledge", "1. Oral questioning"]},
        ],
        "review": ["Trainer:", "Summarizes."],
        "assignment": "Build a questionnaire.",
        "lln_requirements": "Provide simplified handouts.",
        "safety_requirements": "Observe lab SOPs.",
    }
    body = ai_client._merge_session_plan_body(ai_rows, _unit(), _session(), 120)
    assert len(body["delivery_steps"]) == 1
    assert body["delivery_steps"][0].minutes == 40
    assert body["assignment"] == "Build a questionnaire."
    assert body["lln_requirements"] == "Provide simplified handouts."


def test_merge_drops_blank_steps_then_falls_back():
    # a step with no trainer/trainee content is dropped; with none left we fall back
    ai_rows = {"delivery_steps": [{"step_label": "Step 1", "minutes": 10,
                                   "trainer_activity": [], "trainee_activity": [],
                                   "learning_check": []}]}
    body = ai_client._merge_session_plan_body(ai_rows, _unit(), _session(), 120)
    assert body["delivery_steps"]                  # fallback kicked in


def test_merge_coerces_string_to_single_line_no_charsplit():
    # a bare string must not be iterated character-by-character
    ai_rows = {"introduction": "Trainer takes roll call and reviews the session."}
    body = ai_client._merge_session_plan_body(ai_rows, _unit(), _session(), 120)
    assert body["introduction"] == ["Trainer takes roll call and reviews the session."]


# --------------------------------------------------------------------------- #
# .docx rendering
# --------------------------------------------------------------------------- #
def _shaded_fill(cell):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def test_docx_has_title_grid_and_banners():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)

    # title
    assert doc.paragraphs[0].text == "SESSION PLAN"
    assert doc.paragraphs[0].runs[0].bold is True

    # portrait Letter page
    sec = doc.sections[0]
    assert round(sec.page_width.cm, 1) == 21.6
    assert round(sec.page_height.cm, 1) == 27.9

    table = doc.tables[0]
    assert table.style.name == "Table Grid"
    grid = [round(int(g.get(qn("w:w"))) / 20 / 28.3465, 2)
            for g in table._element.find(qn("w:tblGrid")).findall(qn("w:gridCol"))]
    assert grid == session_plan_builder.GRID_WIDTHS

    # the four grey section banners are present and shaded
    banner_texts = {row.cells[0].paragraphs[0].text for row in table.rows
                    if _shaded_fill(row.cells[0]) == "D9D9D9"}
    assert "Session Presentation" in banner_texts
    assert "2. Session Delivery" in banner_texts
    assert any(t.startswith("3.") for t in banner_texts)


def _cell_texts(doc):
    return ["\n".join(p.text for p in cell.paragraphs)
            for t in doc.tables for row in t.rows for cell in row.cells]


def test_docx_session_presentation_and_total_are_centered():
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)

    def para_for(prefix):
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    p = cell.paragraphs[0]
                    if p.text.strip().startswith(prefix):
                        return p
        return None

    assert para_for("Session Presentation").alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert para_for("TOTAL TIME").alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_docx_intro_and_review_are_bulleted():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)
    text = "\n".join(_cell_texts(doc))
    assert "•  Take roll call." in text                       # imperative + bullet
    assert "•  Summarize the key points covered in the session." in text


def test_docx_reflection_has_no_red_hint():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)
    reflection = next(c for c in _cell_texts(doc)
                      if c.strip().startswith("Session Reflection"))
    assert reflection.strip() == "Session Reflection:"        # no "(remarks...)" hint


def test_activity_lines_strips_repeated_role_prefix():
    lines = ["Trainer:", "Trainer: Take roll call.",
             "Trainer - Review the previous session.", "State the outcomes."]
    assert session_plan_builder._activity_lines(lines) == [
        "Take roll call.", "Review the previous session.", "State the outcomes."]


def test_docx_intro_bullets_do_not_repeat_trainer():
    # AI sometimes prefixes EVERY intro line with 'Trainer:' - it must appear once.
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    plan.introduction = ["Trainer: Take roll call.",
                         "Trainer: Review the previous session.",
                         "Trainer: State the session outcomes."]
    doc = session_plan_builder.build_session_plan_document(plan)
    intro = next(c for c in _cell_texts(doc)
                 if c.strip().startswith("Trainer:") and "roll call" in c)
    # exactly one 'Trainer:' - the bold section label, none on the bullets
    assert intro.count("Trainer:") == 1
    assert "•  Take roll call." in intro


def test_merge_preserves_substep_labels_and_minutes():
    ai_rows = {
        "delivery_steps": [
            {"step_label": "Step 1(a)", "minutes": 30,
             "trainer_activity": ["Trainer:", "Introduce the tool."],
             "trainee_activity": ["Trainee(s):", "Watch and take notes."],
             "learning_check": ["Knowledge", "1. Oral questioning"]},
            {"step_label": "Step 1(b)", "minutes": 40,
             "trainer_activity": ["Trainer:", "Guide the hands-on build."],
             "trainee_activity": ["Trainee(s):", "Build a questionnaire."],
             "learning_check": ["Skills", "1. Observation"]},
            {"step_label": "Step 2", "minutes": 40,
             "trainer_activity": ["Trainer:", "Facilitate presentations."],
             "trainee_activity": ["Trainee(s):", "Present findings."],
             "learning_check": ["Attitudes", "1. Teamwork"]},
        ],
    }
    body = ai_client._merge_session_plan_body(ai_rows, _unit(), _session(), 120)
    assert [s.step_label for s in body["delivery_steps"]] == [
        "Step 1(a)", "Step 1(b)", "Step 2"]
    assert sum(s.minutes for s in body["delivery_steps"]) == 110   # 120 - 10


def test_docx_renders_substep_rows():
    from models import DeliveryStep
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    plan.delivery_steps = [
        DeliveryStep("Step 1(a)", 30, ["Trainer:", "Introduce."],
                     ["Trainee(s):", "Observe."], ["Knowledge", "1. Oral questioning"]),
        DeliveryStep("Step 1(b)", 40, ["Trainer:", "Guide the build."],
                     ["Trainee(s):", "Build."], ["Skills", "1. Observation"]),
        DeliveryStep("Step 2", 40, ["Trainer:", "Facilitate."],
                     ["Trainee(s):", "Present."], ["Attitudes", "1. Teamwork"]),
    ]
    doc = session_plan_builder.build_session_plan_document(plan)
    table = doc.tables[0]
    labels = [r.cells[0].paragraphs[0].text for r in table.rows
              if len(r._tr.findall(qn("w:tc"))) == 5
              and r.cells[0].paragraphs[0].text.startswith("Step")]
    assert labels == ["Step 1(a)", "Step 1(b)", "Step 2"]


def test_docx_delivery_rows_match_steps():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)
    table = doc.tables[0]
    # step rows carry 5 <w:tc> cells (label | minutes | trainer[span2] | trainee | check)
    step_rows = [r for r in table.rows
                 if len(r._tr.findall(qn("w:tc"))) == 5
                 and r.cells[0].paragraphs[0].text.startswith("Step")]
    assert len(step_rows) == len(plan.delivery_steps)


def test_docx_multiword_resource_not_charsplit():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)
    text = "\n".join(p.text for t in doc.tables for row in t.rows
                     for cell in row.cells for p in cell.paragraphs)
    assert "Kotler" in text and "Marketing Management" in text


def test_docx_signature_block_has_date_and_sign_10pt():
    plan = ai_client.build_session_plan_offline(
        _unit(), _session(), _inputs(), duration_minutes=120)
    doc = session_plan_builder.build_session_plan_document(plan)
    rows = {role: next((p for p in doc.paragraphs
                        if p.text.strip().startswith(role)), None)
            for role in ("PREPARED BY", "VERIFIED BY", "APPROVED BY")}
    for role, par in rows.items():
        assert par is not None, f"missing {role} row"
        # DATE + SIGN labels preserved on the same row
        assert "DATE:" in par.text and "SIGN:" in par.text
        # 10 pt throughout; only the labels are bold (leaders are not)
        for run in par.runs:
            assert run.font.size == Pt(10)
        bold_labels = {r.text.strip() for r in par.runs if r.bold}
        assert bold_labels == {f"{role}:", "DATE:", "SIGN:"}
